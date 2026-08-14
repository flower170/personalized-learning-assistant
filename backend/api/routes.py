"""
A3 REST API 路由

统一通过 core/graph.py (LangGraph 编排) 或 core/capabilities.py (能力层) 提供服务。
保持与旧前端兼容的端点路径。
"""
from __future__ import annotations

import json
import logging
import os
from pathlib import Path
from typing import AsyncGenerator, Optional

from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel

from core.graph import run_orchestrator, run_capability_stream
from core import get_capability
from services.cache import cache_service
from core.models.schemas import ProfileIncrementUpdateRequest

logger = logging.getLogger(__name__)

_rag_store: dict[str, str] = {}
_rag_cache: dict[str, str] = {}
_doc_sessions: set[str] = set()

DOC_INTERACTION_RULES = """
# 文档学习交互模式规则（高优先级）

当前会话已上传文档并生成了学习资源，你必须遵守以下规则：

1. 所有回答必须基于已上传文档的知识范围，不要脱离文档自由发挥
2. 严禁重新生成全套资源（讲义/思维导图/练习题等），只做增量补充或局部调整
3. 按用户具体需求响应：指定知识点讲解、补充习题、优化格式、代码示例拓展等
4. 如果用户问的内容超出文档范围，友好提示回到文档知识范围内
5. 全程记住当前学习主题，不随意切换无关内容
"""


def _extract_text_from_file(file_path: str, ext: str) -> str:
    try:
        if ext == "pdf":
            import fitz
            doc = fitz.open(file_path)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text
        elif ext in ("docx", "doc"):
            return _extract_docx_robust(file_path)
        elif ext in ("txt", "md"):
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
    except Exception as e:
        logger.warning(f"文本提取失败 ({ext}): {e}")
    return ""


def _extract_docx_robust(file_path: str) -> str:
    """从 docx 提取文本，兼容缺少 footnotes.xml 等非标准结构（python-docx 会崩溃）"""
    import zipfile
    import re

    # 优先尝试 python-docx（结果更干净），失败则回退到直接解析 XML
    try:
        from docx import Document
        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        text = "\n".join(parts)
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"python-docx 提取失败，改用 XML 解析: {e}")

    # 直接解析 word/document.xml，不依赖严格 schema
    import xml.etree.ElementTree as ET
    with zipfile.ZipFile(file_path) as z:
        if "word/document.xml" not in z.namelist():
            raise ValueError("docx 缺少 word/document.xml")
        xml_bytes = z.read("word/document.xml")

    W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    root = ET.fromstring(xml_bytes)

    def para_text(p) -> str:
        parts = []
        for node in p.iter():
            tag = node.tag
            if tag == W_NS + "t":
                parts.append(node.text or "")
            elif tag == W_NS + "tab":
                parts.append("\t")
            elif tag in (W_NS + "br", W_NS + "cr"):
                parts.append("\n")
        return "".join(parts).strip()

    paragraphs = [para_text(p) for p in root.iter(W_NS + "p")]
    return "\n".join(p for p in paragraphs if p)


router = APIRouter()


# ======================== 辅助函数 ========================


async def sse_stream(content_gen: AsyncGenerator[str, None]) -> AsyncGenerator[str, None]:
    """将文本生成器包装为 SSE 格式"""
    try:
        async for chunk in content_gen:
            lines = chunk.replace("\n", "\\n")
            yield f"data: {lines}\n\n"
        yield "data: [DONE]\n\n"
    except Exception as e:
        logger.error(f"SSE 流异常: {e}")
        yield f"data: [ERROR] {str(e)}\n\n"
        yield "data: [DONE]\n\n"


async def event_stream(capability_gen):
    """将能力事件流包装为 SSE 格式"""
    try:
        async for event in capability_gen:
            yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"
    except Exception as e:
        logger.error(f"事件流异常: {e}")
        yield f"data: {json.dumps({'event': 'error', 'message': str(e)[:200]}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"


# ======================== Pydantic Models ========================


class ChatRequest(BaseModel):
    student_id: str = "anonymous"
    message: str
    session_id: str = ""
    explicit_type: str = ""
    language: str = ""
    temp_file_id: Optional[str] = None


class ResourceRequest(BaseModel):
    student_id: str = "anonymous"
    resource_type: str = "lecture"
    topic: str
    course: str = ""
    additional_info: str = ""
    user_demand: str = ""
    temp_file_id: Optional[str] = None


class DispatchRequest(BaseModel):
    student_id: str = "anonymous"
    topic: str
    course: str = ""
    resource_types: list[str] = ["lecture", "mindmap", "exercise", "reading", "code"]
    user_demand: str = ""
    temp_file_id: Optional[str] = None
    language: str = ""


class ProfileChatInitRequest(BaseModel):
    student_id: str
    name: str = ""
    grade: str = ""
    major: str = ""
    language: str = ""


class ProfileChatSendRequest(BaseModel):
    student_id: str
    session_id: str
    message: str
    language: str = ""


class PlanRequest(BaseModel):
    student_id: str = "anonymous"
    topic: str
    course: str = ""
    goal: str = ""
    total_days: int = 30
    daily_minutes: int = 60
    language: str = ""


class TutorRequest(BaseModel):
    student_id: str = "anonymous"
    question: str
    conversation_history: list[dict] = []


# ======================== 大模型自我介绍 ========================

ASSISTANT_INTRO_SYSTEM_PROMPT = """你是「彩迹熊 AI 学习助手」，一个智能学习平台助手。

学生点击了功能按钮，现在轮到你先向学生打招呼。请用**一句简短、亲切**的话：
1. 以"你好"开头
2. 说明你可以解答问题、提供学习资源
3. 以"请问你想要做什么？"结束

直接输出这一句话即可，不要任何额外内容、分点或解释。

示例：你好，我可以为你解答问题，提供学习资源，请问你想要做什么？
"""

_intro_cache: dict = {"content": "", "ts": 0.0}


@router.get("/api/assistant/intro")
async def assistant_intro():
    """大模型自我介绍 + 能力说明（带 5 分钟缓存）"""
    import time as _t

    now = _t.time()
    if _intro_cache["content"] and now - _intro_cache["ts"] < 300:
        return {"reply": _intro_cache["content"]}
    try:
        from services.llm import llm_service
        reply = await llm_service.simple_chat(
            ASSISTANT_INTRO_SYSTEM_PROMPT,
            "请做自我介绍",
            model="spark-lite",
            temperature=0.7,
        )
        if reply:
            _intro_cache["content"] = reply
            _intro_cache["ts"] = now
            return {"reply": reply}
    except Exception as e:
        logger.exception("[assistant_intro] 生成自我介绍失败")
    return {"reply": "你好！我是彩迹熊 AI 学习助手，可以帮你构建学习画像、生成学习资料、制定学习路径、解答学习问题。有什么需要帮助的吗？😊"}


# ======================== 1. 统一聊天 API ========================


@router.post("/api/chat/send")
async def chat_send(req: ChatRequest):
    """统一聊天入口 — LangGraph 编排 + 可选 RAG"""
    if not req.message.strip():
        raise HTTPException(status_code=400, detail="消息不能为空")

    session_key = req.session_id or req.student_id
    if req.temp_file_id:
        _rag_store[session_key] = req.temp_file_id
        _doc_sessions.add(session_key)
        logger.info(f"[Chat] 文件绑定: {session_key} → {req.temp_file_id}")
    active_file_id = req.temp_file_id or _rag_store.get(session_key, "")

    # 混合对话记忆：加载上下文（早期摘要 + 最近 N 轮）
    from services.context_memory import context_memory
    ctx = await context_memory.load_context(req.student_id, session_key)

    # ✅ 有上传文件时：直接走星火知识库问答（本地文件走文本匹配）
    if active_file_id and req.explicit_type != 'resource':
        from services.rag import kb_client
        try:
            if active_file_id.startswith('local_'):
                from services.rag import score_paragraphs, split_paragraphs
                local_dir = Path(__file__).parent.parent / "static" / "uploads"
                txt_path = local_dir / f"{active_file_id}.txt"
                if txt_path.exists():
                    text = txt_path.read_text(encoding="utf-8")
                    paragraphs = split_paragraphs(text)

                    # 多轮意图跟随：加载历史，用"当前问题+最近用户问题"做检索
                    history_msgs = []
                    history_text = ""
                    try:
                        _ctx = await context_memory.load_context(req.student_id, session_key, recent_n=6)
                        history_msgs = list(_ctx.messages)
                        _summ = str(_ctx.summary or "").strip()
                        history_text = "\n".join(
                            f"{'学生' if m.get('role') == 'user' else '老师'}: {str(m.get('content',''))[:300]}"
                            for m in history_msgs[-6:]
                        )
                        # 早期对话的 LLM 摘要也要带上，避免压缩后长记忆丢失
                        if _summ:
                            history_text = f"[早前对话摘要]\n{_summ}\n\n" + history_text
                    except Exception as e:
                        logger.warning(f"[Chat] 历史加载失败: {e}")

                    retrieval_query = req.message
                    recent_user_qs = [m["content"] for m in history_msgs if m.get("role") == "user"]
                    if recent_user_qs:
                        retrieval_query = req.message + " " + " ".join(recent_user_qs[-3:])

                    scored = score_paragraphs(retrieval_query, paragraphs)
                    want_more = any(k in req.message for k in ("习题", "练习", "作业", "题目"))
                    top_n = 12 if want_more else 3
                    top = [p for s, p in scored][:top_n] or paragraphs[:3]
                    context = "\n\n".join(top) if top else ""
                    reply = ""
                    if context:
                        # 交给大模型整理（历史 + 检索段落 → 自然语言回答）
                        try:
                            from services.llm import llm_service
                            prompt = kb_client.build_rag_prompt(req.message, context)
                            if history_text:
                                prompt = (f"[对话历史]\n{history_text}\n\n---\n\n"
                                          f"[当前问题]\n{req.message}\n\n"
                                          f"[文档参考资料]\n{context}\n\n"
                                          "请结合对话历史理解当前问题的意图（若当前问题指代模糊如"
                                          "\"具体是怎么实现\"\"它\"等，应回到历史最近讨论的主题上），"
                                          "并基于文档参考资料用中文回答学生。回答结构清晰、口语化，适合学生理解。")
                            system = ("你是「彩迹熊 AI 学习助手」。请基于提供的文档参考资料，用中文回答学生的问题。"
                                      "要求：1. 优先基于参考资料回答；2. 引用文档内容时保持原文信息准确；"
                                      "3. 资料不足以完整回答时如实说明；4. 回答结构清晰、口语化，适合学生理解。")
                            reply = await llm_service.simple_chat(system, prompt, model="spark-4.0-ultra")
                            if reply and reply.startswith("\n\n[生成中断"):
                                reply = ""
                        except Exception as e:
                            logger.warning(f"[Chat] 本地检索大模型整理失败，回退原文: {e}")
                    reply = reply or context or "文件内容已提取，请尝试更具体的问题。"
                    await context_memory.append_turn(req.student_id, session_key, req.message, reply)
                    return {"reply": reply, "intent": "tutor", "session_id": req.session_id}
            else:
                reply_parts = []
                async for chunk in kb_client.chat_stream(req.message, active_file_id):
                    reply_parts.append(chunk)
                reply = "".join(reply_parts)
                if reply and not reply.startswith("\n\n[知识库错误"):
                    await context_memory.append_turn(req.student_id, session_key, req.message, reply)
                    return {"reply": reply, "intent": "tutor", "session_id": req.session_id}
        except Exception as e:
            logger.warning(f"[Chat] KB 直答失败: {e}")

    try:
        result = await run_orchestrator(
            user_id=req.student_id,
            message=req.message,
            session_id=req.session_id,
            explicit_type=req.explicit_type,
            language=req.language,
            context_summary=ctx.summary,
            context_history=ctx.messages,
        )
        reply = ""
        if result.get("intent") == "unknown" and len(req.message) > 10:
            logger.info("[Chat] unknown + 长消息 → 改为 tutor")
            from core.capabilities.impl.tutor_agent import tutor_agent
            session_key = req.session_id or req.student_id
            rag_ctx = _rag_cache.get(session_key, "")
            question = req.message
            if session_key in _doc_sessions:
                question = DOC_INTERACTION_RULES + "\n\n" + question
            if rag_ctx:
                question = f"请基于以下文件内容回答问题。\n\n[文件内容]\n{rag_ctx[:3000]}\n\n用户问题：{req.message}"
            try:
                parts = []
                async for chunk in tutor_agent.answer(
                    question, req.student_id,
                    conversation_history=ctx.messages or None,
                    language=req.language,
                    context_summary=ctx.summary,
                ):
                    parts.append(chunk)
                reply = "".join(parts)
                await context_memory.append_turn(req.student_id, session_key, req.message, reply)
                return {"reply": reply, "intent": "tutor", "session_id": req.session_id, "is_completed": False}
            except Exception as e:
                logger.exception(f"[Chat] tutor fallback: {e}")

        if result.get("profile_reply"):
            reply = result["profile_reply"]
        elif result.get("resource_clarification_msg"):
            reply = result["resource_clarification_msg"]
        elif result.get("resource_skip_sync"):
            reply = result.get("sse_buffer", ["请选择要生成的资源类型"])[-1]
        elif result.get("resource_results"):
            types = list(result["resource_results"].keys())
            reply = f"✅ 已生成 {len(types)} 类资源: {', '.join(types)}"
        elif result.get("plan_result"):
            stages = result["plan_result"].get("stages", [])
            reply = f"✅ 已规划 {len(stages)} 个学习阶段"
        elif result.get("tutor_reply"):
            reply = result["tutor_reply"]
        else:
            reply = result.get("sse_buffer", ["好的"])[-1]
        # 清理模型偶尔回显的"助手:"/"学生:"等前缀标签
        for _p in ("助手：", "助手:", "学生：", "学生:", "AI助手：", "AI助手:", "AI:"):
            if reply.startswith(_p):
                reply = reply[len(_p):].lstrip()
                break
        await context_memory.append_turn(req.student_id, session_key, req.message, reply)
        return {
            "reply": reply,
            "intent": result.get("intent", "unknown"),
            "session_id": result.get("session_id", ""),
            "is_completed": result.get("profile_completed", False),
            "resource_types": result.get("resource_types", []),
            "resource_topic": result.get("resource_topic", ""),
        }
    except Exception as e:
        logger.exception("[统一聊天] 异常")
        raise HTTPException(status_code=500, detail=str(e)[:200])


@router.get("/api/chat/stream")
async def chat_stream(student_id: str, message: str, session_id: str = "", explicit_type: str = "", language: str = ""):
    """SSE 流式聊天（含 15s 心跳保活）"""
    import asyncio

    async def event_stream():
        heartbeat_interval = 15
        try:
            # 混合对话记忆：加载上下文
            from services.context_memory import context_memory
            session_key = session_id or student_id
            ctx = await context_memory.load_context(student_id, session_key)
            result = await run_orchestrator(
                user_id=student_id, message=message,
                session_id=session_id, explicit_type=explicit_type,
                language=language,
                context_summary=ctx.summary,
                context_history=ctx.messages,
            )
            reply = ""
            if result.get("profile_reply"):
                reply = result["profile_reply"]
            elif result.get("resource_results"):
                types = list(result["resource_results"].keys())
                reply = f"✅ 已生成 {len(types)} 类资源: {', '.join(types)}"
            elif result.get("plan_result"):
                stages = result["plan_result"].get("stages", [])
                reply = f"✅ 已规划 {len(stages)} 个学习阶段"
            elif result.get("tutor_reply"):
                reply = result["tutor_reply"]
            else:
                reply = result.get("sse_buffer", ["好的"])[-1]
            await context_memory.append_turn(student_id, session_key, message, reply)
            # 模拟分块输出
            import json as _json
            for i in range(0, len(reply), 3):
                chunk = reply[i:i+3]
                yield f"data: {_json.dumps({'chunk': chunk, 'intent': result.get('intent', 'unknown')})}\n\n"
            yield f"data: {_json.dumps({'done': True, 'intent': result.get('intent', 'unknown')})}\n\n"
        except Exception as e:
            logger.exception("[SSE流] 异常")
            yield f"data: {_json.dumps({'error': str(e)[:100]})}\n\n"
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ======================== 2. 画像相关 API ========================


@router.post("/api/profile/chat/init")
async def profile_chat_init(req: ProfileChatInitRequest):
    """初始化画像对话"""
    if not req.student_id:
        raise HTTPException(status_code=400, detail="学号不能为空")
    try:
        cap = get_capability("profile")
        session_id, first_question = await cap.init_chat(
            student_id=req.student_id,
            name=req.name,
            grade=req.grade,
            major=req.major,
            language=req.language,
        )
        return {
            "session_id": session_id,
            "student_id": req.student_id,
            "is_new": True,
            "first_question": first_question,
            "base_info": {"student_id": req.student_id, "name": req.name,
                          "grade": req.grade, "major": req.major},
        }
    except Exception as e:
        logger.exception("画像对话初始化失败")
        raise HTTPException(status_code=500, detail=str(e)[:200])


@router.post("/api/profile/chat/send")
async def profile_chat_send(req: ProfileChatSendRequest):
    """发送画像对话消息"""
    if not req.message.strip():
        raise HTTPException(status_code=400, detail="消息不能为空")
    try:
        cap = get_capability("profile")
        reply, is_completed, profile, radar_scores = await cap.chat(
            student_id=req.student_id,
            session_id=req.session_id,
            message=req.message,
            language=req.language,
        )
        return {
            "session_id": req.session_id,
            "reply": reply,
            "is_completed": is_completed,
            "profile": profile.model_dump() if profile else None,
            "radar_scores": radar_scores.model_dump() if radar_scores else None,
        }
    except Exception as e:
        logger.exception("画像对话发送失败")
        return {
            "session_id": req.session_id,
            "reply": "抱歉，我刚刚没听清，能再说一遍吗？",
            "is_completed": False,
            "profile": None,
            "radar_scores": None,
        }


@router.get("/api/profile/{student_id}")
async def get_profile(student_id: str):
    """获取学生完整画像"""
    from core.models.profile import profile_manager
    profile = profile_manager.get_profile(student_id)
    return profile.model_dump()


@router.get("/api/profile/{student_id}/radar")
async def get_profile_radar(student_id: str):
    """获取雷达图维度分值"""
    from core.models.profile import profile_manager
    radar = profile_manager.extract_radar_scores(student_id)
    return radar.model_dump()


@router.get("/api/profile/{student_id}/progress")
async def get_profile_progress(student_id: str):
    """获取画像采集进度"""
    from core.capabilities.impl.profile_chat_agent import profile_chat_agent
    session_id = profile_chat_agent.get_session_id(student_id)
    if not session_id:
        return {
            "stage": "idle", "current_dim": None, "dims_done": [],
            "total_dims": 6, "progress_percent": 0, "asked_count": 0,
        }
    progress = profile_chat_agent.get_progress(student_id, session_id)
    return progress.model_dump()


@router.post("/api/profile/reset")
async def reset_profile(student_id: str = "anonymous"):
    """重置学生画像和对话历史"""
    from core.capabilities.impl.profile_chat_agent import profile_chat_agent
    profile_chat_agent.reset_session(student_id)
    return {"status": "ok", "message": f"画像和对话已重置: {student_id}"}


@router.post("/api/profile/update_increment")
async def update_profile_increment(req: ProfileIncrementUpdateRequest):
    """增量更新画像（学习路径模块联动修改学情用）"""
    from core.models.profile import profile_manager
    profile = profile_manager.incremental_update(req.student_id, req.updates)
    radar = profile_manager.extract_radar_scores(req.student_id)
    return {"profile": profile.model_dump(), "radar_scores": radar.model_dump()}


# ======================== 3. 资源生成 API ========================


@router.post("/api/resource/generate")
async def generate_resource(req: ResourceRequest):
    """生成单种学习资源（SSE 流式）"""
    logger.info(f"资源生成请求: type={req.resource_type}, topic={req.topic}")

    async def stream():
        async for event in run_capability_stream(
            "resource",
            student_id=req.student_id,
            topic=req.topic,
            course=req.course,
            resource_types=[req.resource_type],
            additional_info=req.additional_info,
            user_demand=req.user_demand,
            temp_file_id=req.temp_file_id,
        ):
            yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/api/dispatch/generate")
async def dispatch_generate(req: DispatchRequest):
    """多智能体协同调度（SSE 流式，异常兜底）"""
    async def event_stream():
        try:
            async for event in run_capability_stream(
                "resource",
                student_id=req.student_id,
                topic=req.topic,
                course=req.course,
                resource_types=req.resource_types,
                user_demand=req.user_demand,
                temp_file_id=req.temp_file_id,
                language=req.language,
            ):
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.exception("[dispatch_generate] 生成异常")
            err = {"event": "error", "message": f"生成失败: {str(e)[:150]}"}
            yield f"data: {json.dumps(err, ensure_ascii=False)}\n\n"
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/api/dispatch/generate-all")
async def dispatch_generate_all(req: DispatchRequest):
    """生成全部5种资源（SSE 流式）"""
    req.resource_types = ["lecture", "mindmap", "exercise", "reading", "code"]
    return await dispatch_generate(req)


# ======================== 4. 学习路径 API ========================


@router.post("/api/plan/generate")
async def plan_generate(req: PlanRequest):
    """生成学习路径"""
    if not req.topic.strip():
        raise HTTPException(status_code=400, detail="知识点不能为空")
    try:
        cap = get_capability("plan")
        plan = await cap.generate_plan(
            student_id=req.student_id,
            topic=req.topic,
            course=req.course,
            goal=req.goal,
            total_days=req.total_days,
            daily_minutes=req.daily_minutes,
            language=req.language,
        )
        return plan
    except Exception as e:
        logger.exception("[路径规划] 异常")
        raise HTTPException(status_code=500, detail=str(e)[:200])


@router.post("/api/learning-path/generate")
async def generate_learning_path(req: PlanRequest):
    """初始化学习路径（兼容旧接口）"""
    from core.capabilities.impl.learning_path_agent import learning_path_agent
    from core.models.learning_path_data import save_path
    path = await learning_path_agent.generate_path(
        req.topic, req.course, req.student_id, "", req.goal,
        req.total_days, req.daily_minutes, "balanced",
    )
    save_path(req.student_id, path)
    return path


@router.get("/api/learning-path/{student_id}")
async def get_learning_path(student_id: str):
    """获取当前学习路径及进度"""
    from core.models.learning_path_data import load_path, calc_progress
    data = load_path(student_id)
    if not data:
        raise HTTPException(status_code=404, detail="未找到学习路径")
    prog = calc_progress(data)
    return {"path": data, "progress": prog}


# ======================== 5. 智能辅导 API ========================


@router.post("/api/tutor/ask")
async def tutor_ask(req: TutorRequest):
    """智能辅导：流式答疑"""
    if not req.question.strip():
        raise HTTPException(status_code=400, detail="问题不能为空")

    from core.capabilities.impl.tutor_agent import tutor_agent

    async def stream():
        # 混合对话记忆：加载上下文并合并客户端历史
        from services.context_memory import context_memory
        ctx = await context_memory.load_context(req.student_id, req.student_id)
        merged = (ctx.messages or []) + (req.conversation_history or [])
        async for chunk in tutor_agent.answer(
            question=req.question,
            student_id=req.student_id,
            conversation_history=merged[-6:],
            context_summary=ctx.summary,
        ):
            yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ======================== 6. 练习交互 API ========================


class ExerciseModifyRequest(BaseModel):
    student_id: str = "anonymous"
    topic: str
    exercises: str  # JSON string of current exercises
    feedback: str = ""


class ExerciseSummaryRequest(BaseModel):
    student_id: str = "anonymous"
    topic: str
    exercises: str  # JSON string of exercises
    answers: str = ""  # JSON string of user answers
    language: str = ""


@router.post("/api/exercise/modify")
async def exercise_modify(req: ExerciseModifyRequest):
    """交互式修改练习题"""
    from core.capabilities.impl.resource_agents import exercise_agent

    try:
        exercises_list = json.loads(req.exercises) if isinstance(req.exercises, str) else req.exercises
    except Exception:
        exercises_list = []

    async def stream():
        feedback = req.feedback or "请调整这些题目的难度和类型，让它们更适合我的水平"
        async for chunk in exercise_agent.modify_exercises(
            original_exercises=exercises_list,
            feedback=feedback,
            student_id=req.student_id,
        ):
            yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/api/exercise/summarize")
async def exercise_summarize(req: ExerciseSummaryRequest):
    """总结练习情况并给出下一步建议"""
    async def stream():
        try:
            exercises_data = json.loads(req.exercises) if isinstance(req.exercises, str) else req.exercises
            answers_data = json.loads(req.answers) if isinstance(req.answers, str) else {}
        except Exception:
            exercises_data = []
            answers_data = {}

        from core.capabilities.impl.tutor_agent import tutor_agent

        # ---- 服务端客观重判：不信任前端 correct，也不让 LLM 从原始 JSON 数个数（曾把错 9 题报成正确率 90%）----
        from core.models.practice_data import _judge_answer as _judge

        total_n = len(exercises_data)
        answered = correct = wrong = pending = 0
        verdict_lines = []
        for idx, ex in enumerate(exercises_data):
            qid = ex.get("id") or f"ex_{idx}"
            ua = None
            if isinstance(answers_data, dict):
                rec = answers_data.get(qid)
                if isinstance(rec, dict):
                    ua = rec.get("userAnswer")
            q_short = (ex.get("question") or "").strip().replace("\n", " ")
            if len(q_short) > 26:
                q_short = q_short[:26] + "…"
            if ua is None or str(ua).strip() == "":
                pending += 1
                verdict_lines.append(f"- 题{idx + 1}：未作答 — {q_short}")
                continue
            answered += 1
            res = _judge(ex.get("type"), ex.get("answer"), ua)
            if res is True:
                correct += 1
                verdict_lines.append(f"- 题{idx + 1}：✓ 正确 — {q_short}（你答：{str(ua)[:40]}）")
            elif res is False:
                wrong += 1
                verdict_lines.append(
                    f"- 题{idx + 1}：✗ 错误 — {q_short}（你答：{str(ua)[:40]}；正确答案：{ex.get('answer')}）")
            else:
                pending += 1
                verdict_lines.append(f"- 题{idx + 1}：⚠ 简答/主观题待人工批改 — {q_short}（你答：{str(ua)[:40]}）")
        accuracy = round(100 * correct / answered) if answered else 0

        # 总体情况由系统确定性生成（LLM 从原始 JSON 数个数不可靠，曾把错 9 题报成正确率 90%）
        fixed_prefix = (
            "## 📊 练习总结\n\n"
            "### 总体情况\n"
            f"- 总题数：{total_n} 题\n"
            f"- 已答题数：{answered} 题\n"
            f"- 正确：{correct} 题，错误：{wrong} 题"
            + (f"（未答/待批改 {pending} 题）" if pending else "")
            + "\n"
            f"- 正确率：{accuracy}%\n"
        )

        # 只给模型「逐题判定」与错题内容，不给原始 answers JSON —— 避免它重新数错
        verdict_block = "\n".join(verdict_lines) if verdict_lines else "- （无作答记录）"
        prompt_parts = [
            "你是一个专业的 AI 辅导老师。学生的练习已经由系统逐题判分，请基于「逐题判定」写出针对性的辅导分析。",
        ]
        if req.language:
            lang_map = {"zh-CN": "请使用简体中文回复", "en-US": "请使用英语回复"}
            prompt_parts.append(lang_map.get(req.language, "请使用中文回复"))

        prompt_parts.append(f"\n## 学习主题\n{req.topic}")
        prompt_parts.append(f"\n## 逐题判定\n{verdict_block}")
        prompt_parts.append(f"""
        \n## 输出要求
        只输出以下三个小节，纯文字 Markdown，**禁止使用 Mermaid 图表、流程图、代码块**（不要画 graph TD 等图），也不要重复输出「练习总结 / 总体情况」标题（那部分已生成）：

        ### 薄弱知识点分析
        - 仅针对「逐题判定」中标为 ✗ 错误的题，列出其知识点并分析原因；没有错题就如实说明全部做对，不要硬编

        ### 学习建议
        - 针对薄弱点的专项学习建议
        - 推荐的学习资源类型

        ### 下一步计划
        - 具体可执行的下一步学习行动
        """)

        user_prompt = "\n".join(prompt_parts)
        # 先输出系统判定的总体情况（数字绝对准确），再流式输出模型的分析
        yield f"data: {json.dumps({'chunk': fixed_prefix}, ensure_ascii=False)}\n\n"
        async for chunk in tutor_agent.answer(
            question=user_prompt,
            student_id=req.student_id,
        ):
            yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ======================== 7. 文件上传 API ========================


@router.post("/api/file/upload")
async def upload_file(file: UploadFile = File(...)):
    """上传文件到知识库"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")

    from api.main import _static_dir as STATIC_DIR_BASE
    import uuid, aiofiles, re, time
    from services.rag import kb_client

    UPLOAD_DIR = STATIC_DIR_BASE / "uploads"
    ALLOWED_EXTENSIONS = {"pdf", "doc", "docx", "md", "txt"}

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: .{ext}")

    safe_filename = re.sub(r'[^\w一-鿿.\-]', '_', file.filename)
    local_name = f"{uuid.uuid4().hex[:16]}_{safe_filename}"
    local_path = UPLOAD_DIR / local_name

    content = await file.read()
    file_size = len(content)
    async with aiofiles.open(str(local_path), "wb") as f:
        await f.write(content)

    file_id = None

    # 方案 A: 上传到讯飞知识库
    try:
        kb_result = await kb_client.upload_document(str(local_path))
        file_id = kb_result.get("data", {}).get("fileId", "") if kb_result.get("code") == 0 else None
        if file_id:
            logger.info(f"讯飞KB上传成功: fileId={file_id}, name={safe_filename}")
        else:
            logger.warning(f"讯飞KB上传返回异常: {json.dumps(kb_result, ensure_ascii=False)[:200]}")
    except Exception as e:
        logger.warning(f"讯飞KB上传失败: {e}")

    # 方案 B: 降级为本地文本提取
    if not file_id:
        try:
            text = _extract_text_from_file(str(local_path), ext)
            if text:
                local_id = f"local_{uuid.uuid4().hex[:12]}"
                local_txt_path = UPLOAD_DIR / f"{local_id}.txt"
                async with aiofiles.open(str(local_txt_path), "w", encoding="utf-8") as f:
                    await f.write(text)
                file_id = local_id
                logger.info(f"本地文本提取成功: {len(text)} 字符, id={local_id}")
        except Exception as e:
            logger.warning(f"本地文本提取也失败: {e}")

    return {
        "success": bool(file_id),
        "temp_file_id": file_id,
        "filename": file.filename,
        "size": file_size,
        "fallback": file_id and file_id.startswith("local_"),
    }


@router.post("/api/kb/chat")
async def kb_chat(req: ChatRequest):
    """知识库文档直接问答（流式 SSE）"""
    if not req.message.strip():
        raise HTTPException(status_code=400, detail="消息不能为空")
    if not req.temp_file_id:
        raise HTTPException(status_code=400, detail="需要文件ID")

    from services.rag import kb_client
    from api.main import _static_dir as STATIC_DIR_BASE

    async def stream():
        try:
            import asyncio as _aio_kb
            await _aio_kb.sleep(0)
            yield f"data: {json.dumps({'event': 'handshake', 'stage': 'start', 'temp_file_id': req.temp_file_id, 'is_local': req.temp_file_id.startswith('local_')}, ensure_ascii=False)}\n\n"
            if req.temp_file_id.startswith('local_'):
                # 本地文件：检索相关段落 → 大模型整理回答（LLM 失败则回退原文段落）
                from services.rag import score_paragraphs, split_paragraphs, kb_client
                txt_path = os.path.join(str(STATIC_DIR_BASE), "uploads", f"{req.temp_file_id}.txt")
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8") as fh:
                        text = fh.read()
                    paragraphs = split_paragraphs(text)

                    # 多轮意图跟随：加载对话历史，把历史作为检索与回答的上下文
                    session_key = req.session_id or req.student_id
                    history_msgs = []
                    history_text = ""
                    try:
                        from services.context_memory import context_memory
                        _ctx = await context_memory.load_context(req.student_id, session_key, recent_n=6)
                        history_msgs = list(_ctx.messages)
                        _summ = str(_ctx.summary or "").strip()
                        history_text = "\n".join(
                            f"{'学生' if m.get('role') == 'user' else '老师'}: {str(m.get('content',''))[:300]}"
                            for m in history_msgs[-6:]
                        )
                        # 早期对话的 LLM 摘要也要带上，避免压缩后长记忆丢失
                        if _summ:
                            history_text = f"[早前对话摘要]\n{_summ}\n\n" + history_text
                    except Exception as e:
                        logger.warning(f"[KB Chat] 历史加载失败: {e}")

                    # 检索查询 = 当前问题 + 最近用户问题（让"具体是怎么实现"关联上文"第五章"）
                    retrieval_query = req.message
                    recent_user_qs = [m["content"] for m in history_msgs if m.get("role") == "user"]
                    if recent_user_qs:
                        retrieval_query = req.message + " " + " ".join(recent_user_qs[-3:])

                    scored = score_paragraphs(retrieval_query, paragraphs)
                    # 习题类问题需要更多召回（练习 5.1~5.13 等多条）
                    want_more = any(k in req.message for k in ("习题", "练习", "作业", "题目"))
                    top_n = 12 if want_more else 5
                    top = [p for s, p in scored][:top_n] or paragraphs[:3]
                    context = "\n\n".join(top)
                    # 交给大模型整理（历史 + 参考内容 + 问题 → 自然语言回答）
                    context_used = False
                    try:
                        from services.llm import llm_service
                        prompt = kb_client.build_rag_prompt(req.message, context)
                        if history_text:
                            prompt = (f"[对话历史]\n{history_text}\n\n---\n\n"
                                      f"[当前问题]\n{req.message}\n\n"
                                      f"[文档参考资料]\n{context}\n\n"
                                      "请结合对话历史理解当前问题的意图（若当前问题指代模糊如"
                                      "\"具体是怎么实现\"\"它\"等，应回到历史最近讨论的主题上），"
                                      "并基于文档参考资料用中文回答学生。回答结构清晰、口语化，适合学生理解。")
                        system = ("你是「彩迹熊 AI 学习助手」。请基于提供的文档参考资料，用中文回答学生的问题。"
                                  "要求：1. 优先基于参考资料回答；2. 引用文档内容时保持原文信息准确；"
                                  "3. 资料不足以完整回答时如实说明；4. 回答结构清晰、口语化，适合学生理解。")
                        async for chunk in llm_service.simple_chat_stream(system, prompt, model="spark-4.0-ultra"):
                            if chunk and not chunk.startswith("\n\n[生成中断"):
                                yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
                                context_used = True
                    except Exception as e:
                        logger.warning(f"[KB Chat] 本地检索大模型整理失败，回退原文: {e}")
                    # LLM 无输出或失败时，回退为原文段落（保证不空回复）
                    if not context_used:
                        for para in top:
                            for i in range(0, len(para), 50):
                                yield f"data: {json.dumps({'chunk': para[i:i+50]}, ensure_ascii=False)}\n\n"
                                import asyncio
                                await asyncio.sleep(0.01)
                    # 记录本轮对话（用于后续轮次的意图跟随）
                    try:
                        from services.context_memory import context_memory
                        await context_memory.append_turn(
                            req.student_id, session_key,
                            req.message,
                            "\n".join(top)[:500] or "（无有效内容）",
                        )
                    except Exception as e:
                        logger.warning(f"[KB Chat] 历史记录失败: {e}")
                else:
                    yield f"data: {json.dumps({'chunk': '文件内容未找到，请重新上传。'}, ensure_ascii=False)}\n\n"
            else:
                async for chunk in kb_client.chat_stream(
                    question=req.message,
                    file_id=req.temp_file_id,
                ):
                    yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.exception(f"[KB Chat] 异常: {e}")
            yield f"data: {json.dumps({'error': str(e)[:200]}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ======================== 8. 系统信息 API ========================


@router.get("/api/models")
async def list_models():
    """列出可用模型"""
    from core.config import settings
    return {
        "models": [
            {"name": m.name, "domain": m.domain, "url": m.api_url}
            for m in settings.MODEL_MAP.values()
        ],
    }


@router.get("/api/capabilities")
async def list_capabilities():
    """列出系统能力清单"""
    from core import capability_registry, tool_registry
    return {
        "capabilities": [
            {"name": m["name"], "description": m["description"],
             "stages": m["stages"], "tools_used": m["tools_used"]}
            for m in capability_registry.list_capabilities()
        ],
        "tools": [
            {"name": t["name"], "description": t["description"]}
            for t in tool_registry.list_tools()
        ],
    }
